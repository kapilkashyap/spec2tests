"""Document text extraction service.

Provides format-aware extraction of plain text from uploaded specification
documents (PDF, DOCX, TXT). The extracted, normalised text is what downstream
services (e.g. the Gemini-powered test-case generator) consume.

Design notes
------------
- Extraction is performed entirely in-memory (via ``io.BytesIO``); no
  temporary files are written to disk.
- Each supported format has a dedicated ``_extract_*`` function that raises
  a domain-specific :class:`DocumentExtractionError` subclass on failure,
  which the API layer maps to appropriate HTTP responses.
- Text is normalised (line-ending unification, trailing whitespace removal,
  collapsing of excessive blank lines) so that downstream consumers receive
  consistent, predictable input.
"""

from __future__ import annotations

import io
import logging
import re

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import DependencyError, EmptyFileError, PdfReadError

from app.config import get_settings
from app.models.document import ExtractedDocument

logger = logging.getLogger(__name__)

# MIME types recognised for each supported extension. Used to populate the
# `content_type` field of the extraction response regardless of what the
# client reported in the multipart upload (which cannot always be trusted).
_CONTENT_TYPES: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}

# Text encodings attempted (in order) when decoding plain-text uploads.
_TEXT_ENCODINGS: tuple[str, ...] = ("utf-8", "utf-8-sig", "utf-16", "latin-1")

# Collapse 3+ consecutive newlines (with optional whitespace between them)
# down to a single blank line for readability and consistent downstream
# parsing.
_MULTI_BLANK_LINE_RE = re.compile(r"\n[ \t]*\n(?:[ \t]*\n)+")

# Collapse runs of horizontal whitespace (but not newlines) into a single space.
_INTRA_LINE_WHITESPACE_RE = re.compile(r"[ \t\u00a0]+")


class DocumentExtractionError(Exception):
    """Base class for all document extraction failures.

    Attributes:
        error_code: Stable, machine-readable identifier for the failure
            reason, intended for use in API error responses.
    """

    error_code: str = "extraction_failed"

    def __init__(self, message: str) -> None:
        super().__init__(message)
        self.message = message


class UnsupportedFileTypeError(DocumentExtractionError):
    """Raised when the uploaded file's extension is not supported."""

    error_code = "unsupported_file_type"


class EmptyFileUploadError(DocumentExtractionError):
    """Raised when the uploaded file has no content."""

    error_code = "empty_file"


class FileTooLargeError(DocumentExtractionError):
    """Raised when the uploaded file exceeds the configured size limit."""

    error_code = "file_too_large"


class CorruptedDocumentError(DocumentExtractionError):
    """Raised when the document cannot be parsed because it is malformed."""

    error_code = "corrupted_document"


class EncryptedDocumentError(DocumentExtractionError):
    """Raised when the document is password-protected and cannot be read."""

    error_code = "encrypted_document"


class NoExtractableTextError(DocumentExtractionError):
    """Raised when parsing succeeds but no meaningful text content is found."""

    error_code = "no_extractable_text"


def _get_extension(filename: str) -> str:
    """Return the lower-cased file extension (including the leading dot).

    Raises:
        UnsupportedFileTypeError: If ``filename`` has no extension.
    """
    if not filename or "." not in filename:
        raise UnsupportedFileTypeError(
            f"Could not determine file type for '{filename}': missing file extension."
        )
    extension = "." + filename.rsplit(".", 1)[-1].strip().lower()
    return extension


def _normalize_text(raw_text: str) -> str:
    """Normalise whitespace in extracted text for consistent downstream use.

    - Unifies CRLF/CR line endings to LF.
    - Strips trailing whitespace from each line.
    - Collapses runs of horizontal whitespace into single spaces.
    - Collapses 3+ consecutive blank lines into a single blank line.
    - Strips leading/trailing blank lines from the document as a whole.
    """
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [
        _INTRA_LINE_WHITESPACE_RE.sub(" ", line).rstrip() for line in text.split("\n")
    ]
    text = "\n".join(lines)
    text = _MULTI_BLANK_LINE_RE.sub("\n\n", text)
    return text.strip()


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int, list[str]]:
    """Extract text from a PDF file's raw bytes.

    Returns:
        A tuple of ``(text, page_count, warnings)``.

    Raises:
        EncryptedDocumentError: If the PDF is password-protected and cannot
            be decrypted with an empty password.
        CorruptedDocumentError: If the PDF cannot be parsed at all.
    """
    try:
        reader = PdfReader(io.BytesIO(file_bytes), strict=False)
    except (EmptyFileError, PdfReadError, DependencyError) as exc:
        raise CorruptedDocumentError(
            "The uploaded PDF file could not be read; it may be corrupted or malformed."
        ) from exc
    except Exception as exc:  # pragma: no cover - defensive catch-all
        raise CorruptedDocumentError(
            "The uploaded PDF file could not be read; it may be corrupted or malformed."
        ) from exc

    if reader.is_encrypted:
        try:
            # Some PDFs are "encrypted" with an empty user password purely to
            # set permissions; attempt a blank-password decrypt before giving up.
            result = reader.decrypt("")
        except Exception as exc:
            raise EncryptedDocumentError(
                "The uploaded PDF is password-protected and cannot be processed."
            ) from exc
        if not result:
            raise EncryptedDocumentError(
                "The uploaded PDF is password-protected and cannot be processed."
            )

    warnings: list[str] = []
    page_texts: list[str] = []
    page_count = len(reader.pages)

    for index, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - defensive per-page guard
            logger.warning("Failed to extract text from PDF page %d: %s", index + 1, exc)
            warnings.append(f"Page {index + 1} could not be read and was skipped.")
            continue
        if not page_text.strip():
            warnings.append(f"Page {index + 1} contained no extractable text.")
        page_texts.append(page_text)

    return "\n\n".join(page_texts), page_count, warnings


def _extract_docx_text(file_bytes: bytes) -> tuple[str, int, list[str]]:
    """Extract text from a DOCX file's raw bytes, preserving reading order.

    Iterates the document body in document order so that paragraphs and
    tables are interleaved as they appear in the source file, and renders
    table rows as pipe-delimited lines to retain tabular structure.

    Returns:
        A tuple of ``(text, non_empty_paragraph_count, warnings)``.

    Raises:
        CorruptedDocumentError: If the DOCX package cannot be opened/parsed.
    """
    try:
        document = DocxDocument(io.BytesIO(file_bytes))
    except PackageNotFoundError as exc:
        raise CorruptedDocumentError(
            "The uploaded DOCX file could not be read; it may be corrupted or is not a "
            "valid Word document."
        ) from exc
    except Exception as exc:  # pragma: no cover - defensive catch-all
        raise CorruptedDocumentError(
            "The uploaded DOCX file could not be read; it may be corrupted or is not a "
            "valid Word document."
        ) from exc

    warnings: list[str] = []
    blocks: list[str] = []
    non_empty_paragraph_count = 0

    try:
        inner_content = list(document.iter_inner_content())
    except Exception as exc:  # pragma: no cover - defensive fallback
        logger.warning(
            "Falling back to sequential paragraph/table extraction for DOCX: %s", exc
        )
        inner_content = list(document.paragraphs) + list(document.tables)

    for item in inner_content:
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text:
                blocks.append(text)
                non_empty_paragraph_count += 1
        elif isinstance(item, Table):
            table_lines = []
            for row in item.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
            if table_lines:
                blocks.append("\n".join(table_lines))
            else:
                warnings.append("An empty table was skipped during extraction.")

    return "\n\n".join(blocks), non_empty_paragraph_count, warnings


def _extract_txt_text(file_bytes: bytes) -> tuple[str, list[str]]:
    """Decode a plain-text file's raw bytes, trying a series of encodings.

    Returns:
        A tuple of ``(text, warnings)``.

    Raises:
        CorruptedDocumentError: If no supported encoding can decode the bytes.
    """
    warnings: list[str] = []
    for encoding in _TEXT_ENCODINGS:
        try:
            return file_bytes.decode(encoding), warnings
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Last resort: decode leniently, replacing invalid sequences, so the
    # caller at least gets *something* usable rather than an outright failure.
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        warnings.append(
            "The file's text encoding could not be determined; some characters may "
            "have been replaced."
        )
        return text, warnings
    except Exception as exc:  # pragma: no cover - decode with errors="replace" cannot fail
        raise CorruptedDocumentError(
            "The uploaded text file could not be decoded with any supported encoding."
        ) from exc


def extract_document(
    filename: str, file_bytes: bytes, declared_content_type: str | None = None
) -> ExtractedDocument:
    """Extract normalised plain text and metadata from an uploaded document.

    Args:
        filename: The original filename provided by the client, used to
            determine the document format via its extension.
        file_bytes: The raw file content.
        declared_content_type: The MIME type reported by the client, if any
            (informational only; extraction relies on the file extension and
            actual byte content, not this value).

    Returns:
        An :class:`ExtractedDocument` describing the extracted content.

    Raises:
        UnsupportedFileTypeError: If the file extension is not one of the
            configured ``allowed_upload_extensions``.
        EmptyFileUploadError: If ``file_bytes`` is empty.
        FileTooLargeError: If ``file_bytes`` exceeds the configured maximum
            upload size.
        CorruptedDocumentError: If the file cannot be parsed as a valid
            document of its declared type.
        EncryptedDocumentError: If a PDF document is password-protected.
        NoExtractableTextError: If parsing succeeds but yields no usable text.
    """
    settings = get_settings()

    if not file_bytes:
        raise EmptyFileUploadError(f"The uploaded file '{filename}' is empty.")

    if len(file_bytes) > settings.max_upload_size_bytes:
        raise FileTooLargeError(
            f"The uploaded file '{filename}' ({len(file_bytes)} bytes) exceeds the maximum "
            f"allowed size of {settings.max_upload_size_mb} MB."
        )

    extension = _get_extension(filename)
    if extension not in settings.allowed_upload_extensions:
        allowed = ", ".join(settings.allowed_upload_extensions)
        raise UnsupportedFileTypeError(
            f"Unsupported file type: '{extension}'. Allowed types: {allowed}"
        )

    page_count: int | None = None
    paragraph_count: int | None = None
    warnings: list[str] = []

    if extension == ".pdf":
        raw_text, page_count, warnings = _extract_pdf_text(file_bytes)
    elif extension == ".docx":
        raw_text, paragraph_count, warnings = _extract_docx_text(file_bytes)
    elif extension == ".txt":
        raw_text, warnings = _extract_txt_text(file_bytes)
    else:  # pragma: no cover - guarded by the allowlist check above
        raise UnsupportedFileTypeError(f"Unsupported file type: '{extension}'.")

    normalized_text = _normalize_text(raw_text)

    if not normalized_text:
        raise NoExtractableTextError(
            f"No extractable text content was found in '{filename}'. The document may be "
            "image-based, empty, or otherwise unreadable."
        )

    content_type = declared_content_type or _CONTENT_TYPES.get(
        extension, "application/octet-stream"
    )

    return ExtractedDocument(
        filename=filename,
        extension=extension,
        content_type=content_type,
        text=normalized_text,
        character_count=len(normalized_text),
        word_count=len(normalized_text.split()),
        page_count=page_count,
        paragraph_count=paragraph_count,
        warnings=warnings,
    )
