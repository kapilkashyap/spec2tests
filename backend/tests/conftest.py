"""Shared pytest fixtures for the Spec2Tests backend test suite.

Provides:
- Deterministic environment isolation: the ``app.config.get_settings``
  LRU cache is cleared before and after every test so that environment
  variable changes made by one test (via ``monkeypatch.setenv``) can never
  leak into another.
- A ready-to-use FastAPI ``TestClient`` bound to the real application
  instance (``app.main.app``), for true end-to-end HTTP-level tests.
- Small in-memory document builders (PDF/DOCX bytes) so extraction tests
  do not depend on any fixture files on disk.
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient


@pytest.fixture(autouse=True)
def _reset_settings_cache():
    """Ensure ``get_settings()`` re-reads the environment for every test.

    Without this, the first test to call ``get_settings()`` would poison
    the ``lru_cache`` for the remainder of the test session, making
    ``monkeypatch.setenv`` calls in later tests silently ineffective.
    """
    from app.config import get_settings

    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture()
def client() -> TestClient:
    """A ``TestClient`` for the real Spec2Tests FastAPI application."""
    from app.main import app

    return TestClient(app)


def _pdf_content_stream(text: str) -> bytes:
    """Build a minimal PDF content stream that draws ``text`` on one line."""
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    return f"BT /F1 18 Tf 72 700 Td ({escaped}) Tj ET".encode("latin-1")


def make_pdf_bytes(pages: list[str]) -> bytes:
    """Build an in-memory, valid, single- or multi-page PDF for tests.

    Args:
        pages: One text string per page; each becomes that page's visible
            (and therefore extractable) content.

    Returns:
        The raw bytes of a minimal but structurally valid PDF document.
    """
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()

    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font)

    for page_text in pages:
        page = writer.add_blank_page(width=612, height=792)

        stream_obj = DecodedStreamObject()
        stream_obj.set_data(_pdf_content_stream(page_text))
        stream_ref = writer._add_object(stream_obj)
        page[NameObject("/Contents")] = stream_ref

        resources = DictionaryObject()
        font_dict = DictionaryObject()
        font_dict[NameObject("/F1")] = font_ref
        resources[NameObject("/Font")] = font_dict
        page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_empty_pdf_bytes() -> bytes:
    """Build a valid, single-page PDF with no text content at all."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_encrypted_pdf_bytes(password: str = "secret123") -> bytes:
    """Build a valid, password-protected PDF for encrypted-document tests."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt(user_password=password)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build an in-memory DOCX document containing paragraphs and an optional table."""
    from docx import Document as DocxDocument

    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row_values in enumerate(table_rows):
            for col_index, value in enumerate(row_values):
                table.rows[row_index].cells[col_index].text = value

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_empty_docx_bytes() -> bytes:
    """Build a valid DOCX document with no paragraph or table content."""
    from docx import Document as DocxDocument

    document = DocxDocument()
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
